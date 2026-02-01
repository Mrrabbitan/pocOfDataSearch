from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, List, Optional

from docx import Document

from ..core.generator import SectionContent


def _clear_document(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)


def _add_bullets(doc: Document, bullets: Iterable[str]) -> None:
    for item in bullets:
        text = str(item).strip()
        if not text:
            continue
        if re.match(r"^(\d+\.|S\d+)", text):
            doc.add_paragraph(text, style="List Number")
        else:
            doc.add_paragraph(text, style="List Bullet")


def _to_chinese_numeral(number: int) -> str:
    numerals = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    if 1 <= number <= 10:
        return numerals[number - 1]
    if 11 <= number <= 19:
        return "十" + numerals[number - 11]
    if number == 20:
        return "二十"
    return str(number)


def _heading_prefix(level: int, indices: List[int]) -> str:
    if level == 1 and indices:
        return f"{_to_chinese_numeral(indices[0])}、"
    if indices:
        return ".".join(str(i) for i in indices)
    return ""


def _render_section(
    doc: Document,
    section: SectionContent,
    level: int,
    indices: List[int],
) -> None:
    prefix = _heading_prefix(level, indices)
    title = f"{prefix} {section.title}".strip()
    doc.add_heading(title, level=min(level, 4))
    for paragraph in section.paragraphs:
        text = str(paragraph).strip()
        if text:
            doc.add_paragraph(text)
    if section.bullets:
        _add_bullets(doc, section.bullets)
    for idx, sub in enumerate(section.subsections, start=1):
        _render_section(doc, sub, level + 1, indices + [idx])


def render_docx(
    sections: List[SectionContent],
    output_path: Path,
    template_path: Optional[Path] = None,
    use_template: bool = False,
) -> Path:
    if use_template and template_path and template_path.exists():
        doc = Document(str(template_path))
        _clear_document(doc)
    else:
        doc = Document()

    for idx, section in enumerate(sections, start=1):
        _render_section(doc, section, level=1, indices=[idx])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
